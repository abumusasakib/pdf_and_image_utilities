import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter.ttk import Progressbar
from io import BytesIO
from PIL import Image
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pytesseract
import threading

# Set the Tesseract executable path
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

stop_requested = False  # Global flag to track stop requests


def convert_to_text_with_ocr(pdf_file, docx_file, progress_label, progress_bar):
    global stop_requested
    stop_requested = False

    pdf_document = fitz.open(pdf_file)
    total_pages = len(pdf_document)
    docx_document = Document()
    progress_label.config(text=f"Starting OCR conversion...")
    progress_bar['maximum'] = total_pages

    for page_number in range(total_pages):
        if stop_requested:
            progress_label.config(text="Conversion stopped by user.")
            return

        progress_label.config(text=f"Processing page {page_number + 1}/{total_pages}...")
        progress_bar['value'] = page_number + 1
        root.update_idletasks()

        page = pdf_document.load_page(page_number)
        pixmap = page.get_pixmap(dpi=300)
        image_bytes = pixmap.tobytes("png")
        image_stream = BytesIO(image_bytes)
        image = Image.open(image_stream)
        extracted_text = pytesseract.image_to_string(image)

        docx_document.add_paragraph(f"Page {page_number + 1}:\n")
        docx_document.add_paragraph(extracted_text)
        if page_number < total_pages - 1:
            docx_document.add_page_break()

    docx_document.save(docx_file)
    progress_label.config(text=f"Conversion completed! Saved as '{docx_file}'.")


def convert_to_centered_images(pdf_file, docx_file, progress_label, progress_bar):
    global stop_requested
    stop_requested = False

    pdf_document = fitz.open(pdf_file)
    total_pages = len(pdf_document)
    docx_document = Document()
    progress_label.config(text=f"Starting conversion...")
    progress_bar['maximum'] = total_pages

    for page_number in range(total_pages):
        if stop_requested:
            progress_label.config(text="Conversion stopped by user.")
            return

        progress_label.config(text=f"Processing page {page_number + 1}/{total_pages}...")
        progress_bar['value'] = page_number + 1
        root.update_idletasks()

        page = pdf_document.load_page(page_number)
        pixmap = page.get_pixmap()
        image_bytes = pixmap.tobytes("png")
        image_stream = BytesIO(image_bytes)
        image = Image.open(image_stream)
        img_width, img_height = image.size
        img_aspect_ratio = img_width / img_height

        max_width = 8.5 - 2 * 0.5
        max_height = 11.0 - 2 * 0.5
        if img_aspect_ratio > 1:
            scaled_width = max_width
            scaled_height = scaled_width / img_aspect_ratio
        else:
            scaled_height = max_height
            scaled_width = scaled_height * img_aspect_ratio

        paragraph = docx_document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        temp_image_stream = BytesIO()
        image.save(temp_image_stream, format="PNG")
        temp_image_stream.seek(0)
        run = paragraph.add_run()
        run.add_picture(temp_image_stream, width=Inches(scaled_width), height=Inches(scaled_height))

        if page_number < total_pages - 1:
            docx_document.add_section(WD_SECTION_START.NEW_PAGE)

    docx_document.save(docx_file)
    messagebox.showinfo("Conversion Completed", f"Conversion completed!\n Saved as '{docx_file}'.")


def start_conversion():
    global stop_requested
    stop_requested = False

    pdf_file = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")])
    if not pdf_file:
        return

    docx_file = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")])
    if not docx_file:
        return

    choice = conversion_choice.get()
    if choice not in ["text", "images"]:
        messagebox.showerror("Error", "Please select a conversion method.")
        return

    def conversion_task():
        if choice == "text":
            convert_to_text_with_ocr(pdf_file, docx_file, progress_label, progress_bar)
        elif choice == "images":
            convert_to_centered_images(pdf_file, docx_file, progress_label, progress_bar)

    threading.Thread(target=conversion_task).start()


def stop_conversion():
    global stop_requested
    stop_requested = True
    progress_label.config(text="Stopping conversion...")


# Tkinter UI
root = tk.Tk()
root.title("PDF to DOCX Converter")

tk.Label(root, text="Select Conversion Method:").pack(pady=10)

conversion_choice = tk.StringVar(value="text")
tk.Radiobutton(root, text="Extract Text with OCR", variable=conversion_choice, value="text").pack()
tk.Radiobutton(root, text="Convert Pages to Centered Images", variable=conversion_choice, value="images").pack()

progress_label = tk.Label(root, text="")
progress_label.pack(pady=10)

progress_bar = Progressbar(root, orient=tk.HORIZONTAL, length=300, mode='determinate')
progress_bar.pack(pady=10)

button_frame = tk.Frame(root)
button_frame.pack(pady=20)

tk.Button(button_frame, text="Start Conversion", command=start_conversion).grid(row=0, column=0, padx=10)
tk.Button(button_frame, text="Stop Conversion", command=stop_conversion).grid(row=0, column=1, padx=10)

root.mainloop()
