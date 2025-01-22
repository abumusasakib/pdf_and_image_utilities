from io import BytesIO
from PIL import Image
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import Tk, filedialog, messagebox


def pdf_to_docx_with_centered_images():
    root = Tk()
    root.withdraw()  # Hide the root window

    try:
        # Ask the user to select a PDF file
        pdf_file = filedialog.askopenfilename(
            title="Select a PDF file",
            filetypes=[("PDF Files", "*.pdf")]
        )
        if not pdf_file:
            messagebox.showinfo("No File Selected", "You must select a PDF file to continue.")
            return

        # Ask the user to specify the output DOCX file
        docx_file = filedialog.asksaveasfilename(
            title="Save as DOCX",
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")]
        )
        if not docx_file:
            messagebox.showinfo("No Save Location", "You must specify a location to save the DOCX file.")
            return

        print(f"\nStarting conversion of {pdf_file} to {docx_file}")

        # Page dimensions and margins
        page_width, page_height = 8.5, 11.0
        margin = 0.5
        max_width = page_width - 2 * margin
        max_height = page_height - 2 * margin

        print("Opening PDF document...")
        pdf_document = fitz.open(pdf_file)
        total_pages = len(pdf_document)
        print(f"Found {total_pages} pages")

        docx_document = Document()

        for page_number in range(total_pages):
            print(f"\nProcessing page {page_number + 1}/{total_pages}")

            page = pdf_document.load_page(page_number)
            print("Converting PDF page to image...")
            pixmap = page.get_pixmap()
            image_bytes = pixmap.tobytes("png")
            image_stream = BytesIO(image_bytes)

            image = Image.open(image_stream)
            img_width, img_height = image.size
            img_aspect_ratio = img_width / img_height
            print(f"Original image dimensions: {img_width}x{img_height} pixels")

            # Scale the image to fit within the page margins
            if img_aspect_ratio > 1:
                scaled_width = max_width
                scaled_height = scaled_width / img_aspect_ratio
            else:
                scaled_height = max_height
                scaled_width = scaled_height * img_aspect_ratio

            if scaled_width > max_width:
                scaled_width = max_width
                scaled_height = scaled_width / img_aspect_ratio
            if scaled_height > max_height:
                scaled_height = max_height
                scaled_width = scaled_height * img_aspect_ratio

            print(f"Scaled dimensions: {scaled_width:.2f}x{scaled_height:.2f} inches")

            # Add the image to the document, centered
            paragraph = docx_document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = paragraph_format.space_after = 0

            print("Adding image to document...")
            temp_image_stream = BytesIO()
            image.save(temp_image_stream, format="PNG")
            temp_image_stream.seek(0)
            run = paragraph.add_run()
            run.add_picture(temp_image_stream, width=Inches(scaled_width), height=Inches(scaled_height))

            # Add a new section for the next page
            if page_number < total_pages - 1:
                docx_document.add_section(WD_SECTION_START.NEW_PAGE)

        print(f"\nSaving document to {docx_file}...")
        docx_document.save(docx_file)
        print("Conversion completed successfully!")
        messagebox.showinfo("Success", f"Conversion completed successfully!\nSaved as: {docx_file}")

    except Exception as e:
        messagebox.showerror("Error", f"An error occurred during conversion:\n{e}")


if __name__ == "__main__":
    pdf_to_docx_with_centered_images()
