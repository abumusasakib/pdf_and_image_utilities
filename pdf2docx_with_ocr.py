from io import BytesIO
from PIL import Image
import fitz  # PyMuPDF
from docx import Document
import pytesseract  # Tesseract OCR
from tkinter import Tk, filedialog, messagebox

# Set the Tesseract executable path
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'


def pdf_to_docx_with_ocr(pdf_file, docx_file):
    # Open the PDF
    pdf_document = fitz.open(pdf_file)
    total_pages = len(pdf_document)

    # Create a new DOCX document
    docx_document = Document()

    print(f"Starting conversion of '{pdf_file}' with {total_pages} pages...")

    # Iterate through each page of the PDF
    for page_number in range(total_pages):
        print(f"Processing page {page_number + 1} of {total_pages}...")

        # Get the page as an image
        page = pdf_document.load_page(page_number)
        pixmap = page.get_pixmap(dpi=300)  # High DPI for better OCR accuracy
        image_bytes = pixmap.tobytes("png")
        image_stream = BytesIO(image_bytes)

        # Open the image with PIL
        image = Image.open(image_stream)

        # Perform OCR on the image
        extracted_text = pytesseract.image_to_string(image)

        # Add the extracted text to the Word document
        docx_document.add_paragraph(f"Page {page_number + 1}:\n")
        docx_document.add_paragraph(extracted_text)

        # Add a page break if there are more pages
        if page_number < total_pages - 1:
            docx_document.add_page_break()

    # Save the DOCX document
    docx_document.save(docx_file)
    print(f"Conversion completed! Saved as '{docx_file}'.")


# GUI for selecting files
def select_files_and_convert():
    root = Tk()
    root.withdraw()  # Hide the main Tkinter window

    try:
        # Ask user to select the PDF file
        pdf_file = filedialog.askopenfilename(
            title="Select a PDF file",
            filetypes=[("PDF Files", "*.pdf")]
        )
        if not pdf_file:
            messagebox.showinfo("No File Selected", "You must select a PDF file to continue.")
            return

        # Ask user to specify the output DOCX file path
        docx_file = filedialog.asksaveasfilename(
            title="Save as DOCX",
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")]
        )
        if not docx_file:
            messagebox.showinfo("No Save Location", "You must specify a location to save the DOCX file.")
            return

        # Perform the conversion
        pdf_to_docx_with_ocr(pdf_file, docx_file)
        messagebox.showinfo("Success", f"Conversion completed successfully!\nSaved as: {docx_file}")

    except Exception as e:
        messagebox.showerror("Error", f"An error occurred during conversion:\n{e}")


if __name__ == "__main__":
    select_files_and_convert()
